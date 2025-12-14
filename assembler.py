import json
import cv2
import os
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def extract_frame_with_redbox(video_path, time_sec, output_path):
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return False
        
    cap.set(cv2.CAP_PROP_POS_MSEC, time_sec * 1000)
    ret1, frame_curr = cap.read()
    
    prev_time = max(0, time_sec - 1.5)
    cap.set(cv2.CAP_PROP_POS_MSEC, prev_time * 1000)
    ret2, frame_prev = cap.read()
    
    cap.release()
    
    if not ret1: return False
    
    final_image = frame_curr.copy()
    
    if ret2:
        try:
            gray_curr = cv2.cvtColor(frame_curr, cv2.COLOR_BGR2GRAY)
            gray_prev = cv2.cvtColor(frame_prev, cv2.COLOR_BGR2GRAY)
            # 1. Analyze Structure (Horizontal Grouping)
            # We want to group [Checkbox] + [Text Label] into one block.
            # So we use a kernel that is wide but short.
            gray_curr_clean = cv2.cvtColor(frame_curr, cv2.COLOR_BGR2GRAY)
            edges = cv2.Canny(gray_curr_clean, 30, 100)
            
            # Horizontal Kernel: (3, 20) means dilate 20 pixels horizontally, 3 vertically
            # This merges words and checkoxes on the same line.
            kernel_horizontal = np.ones((3, 25), np.uint8) 
            edges_dilated = cv2.dilate(edges, kernel_horizontal, iterations=1)
            
            contours_ui, _ = cv2.findContours(edges_dilated, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
            
            ui_boxes = []
            for c in contours_ui:
                x, y, w, h = cv2.boundingRect(c)
                # Filter useful UI rows
                # Height should be reasonable for a text line (15-60px)
                # Width should be enough for text (>50px)
                if h > 15 and h < 80 and w > 50: 
                    ui_boxes.append((x, y, w, h))

            # 2. Analyze Changes (Diff)
            diff = cv2.absdiff(gray_curr, gray_prev)
            _, thresh = cv2.threshold(diff, 25, 255, cv2.THRESH_BINARY)
            
            # Dilate the diff too, to tolerate small discrepancies
            kernel_diff = np.ones((5,5), np.uint8)
            dilated_diff = cv2.dilate(thresh, kernel_diff, iterations=3)
            
            contours_diff, _ = cv2.findContours(dilated_diff, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Sort contours by area
            contours_diff = sorted(contours_diff, key=cv2.contourArea, reverse=True)
            
            count = 1
            boxes_found = []
            
            for c in contours_diff:
                if cv2.contourArea(c) > 300: # Significant change
                    dx, dy, dw, dh = cv2.boundingRect(c)
                    
                    # 3. Smart Match: Find enclosing UI Box
                    best_match = None
                    min_area = float('inf')
                    
                    # Logic: Find smallest UI box that roughly contains the diff box
                    # We allow some tolerance
                    for ux, uy, uw, uh in ui_boxes:
                        # Check intersection/inclusion
                        # Simple inclusion: UI box contains center of Diff box?
                        cx = dx + dw/2
                        cy = dy + dh/2
                        
                        if (ux <= cx <= ux + uw) and (uy <= cy <= uy + uh):
                            # It contains center. Is it smaller than current best?
                            area = uw * uh
                            # Also avoid huge container boxes (like the whole window)
                            if area < min_area and area < (frame_curr.shape[0]*frame_curr.shape[1] * 0.9):
                                min_area = area
                                best_match = (ux, uy, uw, uh)
                    
                    # Select Box to Draw
                    if best_match:
                        fx, fy, fw, fh = best_match
                    else:
                        # Fallback: Just the diff box + padding
                        fx, fy, fw, fh = dx-5, dy-5, dw+10, dh+10
                    
                    # Check overlap with already drawn boxes to prevent duplicates
                    is_duplicate = False
                    for b in boxes_found:
                        # simple centers distance check
                        bx, by, bw, bh = b
                        if abs(fx - bx) < 20 and abs(fy - by) < 20: 
                             is_duplicate = True
                             break
                    
                    if is_duplicate: continue

                    # Draw Red Box (BGR: 0, 0, 255)
                    cv2.rectangle(final_image, (fx, fy), (fx+fw, fy+fh), (0, 0, 255), 2)
                    
                    # Draw Label (Circle with Number)
                    center_x = fx
                    center_y = fy
                    radius = 12
                    # Ensure circle is inside image
                    center_y = max(radius, center_y)
                    center_x = max(radius, center_x)
                    
                    cv2.circle(final_image, (center_x, center_y), radius, (0, 0, 255), -1)
                    
                    label = str(count)
                    cv2.putText(final_image, label, (center_x - 5, center_y + 5), 
                                cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 255), 1, cv2.LINE_AA)
                    
                    boxes_found.append((fx, fy, fw, fh))
                    count += 1
                    if count > 5: break
            
            # Save
            cv2.imwrite(output_path, final_image)
            return len(boxes_found)
            
        except Exception as e:
            print(f"RedBox Error: {e}")
            cv2.imwrite(output_path, final_image)
            return 0
    else:
        cv2.imwrite(output_path, final_image)
        return 0

def set_aifteam_styles(doc):
    # 1. Page Margins
    section = doc.sections[0]
    section.left_margin = Twips(1134)
    section.right_margin = Twips(1418)
    section.top_margin = Twips(1418)
    section.bottom_margin = Twips(1418)

    # 2. Normal Style
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(12)
    # Space After ~10pt (127000 EMU)
    style_normal.paragraph_format.space_after = Pt(10)

    # 3. Heading 1
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Calibri'
    font_h1.size = Pt(14)
    font_h1.color.rgb = RGBColor(0x80, 0x80, 0x80) # Gray
    font_h1.bold = True
    style_h1.paragraph_format.space_before = Pt(12)
    style_h1.paragraph_format.space_after = Pt(0)

    # 4. Heading 2
    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Calibri'
    font_h2.size = Pt(14)
    font_h2.color.rgb = RGBColor(0x80, 0x80, 0x80) # Gray
    font_h2.bold = False
    style_h2.paragraph_format.space_before = Pt(6)
    style_h2.paragraph_format.space_after = Pt(6)

    # 5. Heading 3 (Dark Blue)
    style_h3 = doc.styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Calibri'
    font_h3.size = Pt(12)
    font_h3.color.rgb = RGBColor(0x1F, 0x4D, 0x78) # Dark Blue
    font_h3.bold = True
    style_h3.paragraph_format.space_before = Pt(2)
    style_h3.paragraph_format.space_after = Pt(0)


def clean_and_prepare_template(template_path):
    doc = Document(template_path)
    body = doc.element.body
    
    # 1. Find Start (1. Amaç) and End (First Table)
    start_index = -1
    table_index = -1
    
    history_tbl = None
    if len(doc.tables) > 0:
        history_tbl = doc.tables[0]._element
        try:
            table_index = body.index(history_tbl)
        except ValueError:
            print("Table not found in body")
            
    # Find "1. Amaç"
    for i, child in enumerate(body.iterchildren()):
        if child.tag.endswith('p'):
            text = "".join([t.text for t in child.xpath('.//w:t') if t.text])
            if "1. Amaç" in text:
                start_index = i
                break
                
    if start_index != -1 and table_index != -1:
        print(f"Cleaning content between index {start_index} and {table_index}")
        # Collect elements to remove
        to_remove = []
        for i, child in enumerate(body.iterchildren()):
            if i >= start_index and i < table_index:
                to_remove.append(child)
        
        for child in to_remove:
            body.remove(child)
    else:
        print("Warning: Could not find markers (1. Amaç or Table). Appending to end.")
        
    return doc, history_tbl

def create_doc_from_plan(video_path, plan_path, output_docx):
    # Load Plan
    with open(plan_path, 'r', encoding='utf-8') as f:
        plan = json.load(f)
        
    # Load and Clean Template
    template_file = "Çoklu Para Birimi Sihirbazı Kullanıcı Dokümanı.docx"
    doc, anchor_element = clean_and_prepare_template(template_file)
    
    # Update Title
    for p in doc.paragraphs[:30]:
        if "Çoklu Para Birimi" in p.text or "Sihirbazı" in p.text:
            print(f"Updating Title: {p.text}")
            p.text = "El Terminali WMS Kullanım Kılavuzu"
    
    # Setup Styles (Enforcing calibration just in case)
    set_aifteam_styles(doc) # Re-apply styles
    
    # Update History Table
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # Add a new row for this document
        # Or just update the first row? Let's add a new row
        row = table.add_row()
        row.cells[0].text = '1.0'
        row.cells[1].text = datetime.now().strftime("%d.%m.%Y")
        row.cells[2].text = 'Otomasyon (Jules)'
        row.cells[3].text = 'El Terminali Eğitimi'

    # Content Body
    img_dir = "final_images"
    if not os.path.exists(img_dir):
        os.makedirs(img_dir)
        
    step_counter = 1
    
    # Helper to insert before anchor
    def add_element_before_anchor(element):
        if anchor_element is not None:
             anchor_element.addprevious(element)
    
    for item in plan:
        if item['type'] == 'heading':
            h = doc.add_heading(item['text'], level=item['level'])
            add_element_before_anchor(h._element)
            
        elif item['type'] == 'step':
            # Add Text
            p = doc.add_paragraph(item['text'])
            add_element_before_anchor(p._element)
            
            # Extract and Add Image
            time_sec = item['time']
            img_name = f"step_{step_counter}_{int(time_sec)}.jpg"
            img_path = os.path.join(img_dir, img_name)
            
            print(f"Processing Step {step_counter}: {item['text'][:30]}... at {time_sec}s")
            
            box_count = extract_frame_with_redbox(video_path, time_sec, img_path)
            
            # Update text with references if boxes found
            if box_count and box_count > 0:
                refs = ", ".join([f"Kutu {i+1}" for i in range(box_count)])
                if box_count == 1:
                    p.add_run(f" (Bkz: Kutu 1)")
                else:
                    p.add_run(f" (Bkz: {refs})")

            if box_count is not False: # extract returns count or False
                try:
                    # Adding picture
                    pic_p = doc.add_paragraph()
                    pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = pic_p.add_run()
                    run.add_picture(img_path, width=Inches(6))
                    add_element_before_anchor(pic_p._element)

                    caption = doc.add_paragraph(f"Ekran Görüntüsü {step_counter}")
                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption.runs[0].font.italic = True
                    caption.runs[0].font.size = Pt(10)
                    add_element_before_anchor(caption._element)
                    
                    # Manual Edit Note
                    if box_count == 0:
                        note_p = doc.add_paragraph()
                        note_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run_note = note_p.add_run("*(Lütfen tıklanacak alanı kırmızı kutu ile işaretleyiniz)*")
                        run_note.font.color.rgb = RGBColor(255, 0, 0)
                        run_note.font.bold = True
                        run_note.font.size = Pt(10)
                        add_element_before_anchor(note_p._element)

                except Exception as e:
                    print(f"Error adding image: {e}")
            
            step_counter += 1
            
    doc.save(output_docx)
    print(f"Successfully saved {output_docx}")

if __name__ == "__main__":
    video_file = "1-El Terminali Eğitimi-20250623_092431-Toplantı Kaydı.mp4"
    plan_file = "content_plan.json"
    output_file = "Taslak_Dokuman_v7_Styled.docx"
    
    create_doc_from_plan(video_file, plan_file, output_file)
