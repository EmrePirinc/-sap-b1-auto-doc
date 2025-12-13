from docx import Document
import sys

def read_docx(path):
    try:
        doc = Document(path)
        print(f"--- CONTENT OF {path} ---")
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    # print first 50 non-empty paragraphs to capture style
    doc = Document(r"..\Çoklu Para Birimi Sihirbazı Kullanıcı Dokümanı.docx")
    count = 0
    print("--- SAMPLE TEXT ---")
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and count < 40:
            print(f"{count}: {text}")
            count += 1
