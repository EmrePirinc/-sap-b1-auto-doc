import webvtt
import cv2
import os
import sys
import numpy as np
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load environment variables (Removed)

def time_string_to_seconds(time_str):
    parts = time_str.split(':')
    seconds = 0.0
    if len(parts) == 3:
        seconds += int(parts[0]) * 3600
        seconds += int(parts[1]) * 60
        seconds += float(parts[2])
    elif len(parts) == 2:
        seconds += int(parts[0]) * 60
        seconds += float(parts[1])
    return seconds

def get_frame(cap, time_sec):
    cap.set(cv2.CAP_PROP_POS_MSEC, time_sec * 1000)
    ret, frame = cap.read()
    return ret, frame

def calculate_image_difference(img1_path, img2_path):
    # Returns a percentage difference between two images
    # If images are not same size, returns 100% diff
    img1 = cv2.imread(img1_path)
    img2 = cv2.imread(img2_path)
    
    if img1 is None or img2 is None:
        return 100.0
    
    if img1.shape != img2.shape:
        return 100.0
        
    diff = cv2.absdiff(img1, img2)
    gray_diff = cv2.cvtColor(diff, cv2.COLOR_BGR2GRAY)
    
    # Count non-zero pixels or calculate mean
    score = np.mean(gray_diff)
    return score


# ... (OpenCV parts remain same) ...

def clean_text_with_regex(text):
    # 1. Remove Speaker labels and Timestamps
    text = re.sub(r'Speaker \d+:', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Konuşmacı \d+:', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\d{2}:\d{2}', '', text)
    
    # 2. Aggressive Filler Removal
    fillers = [
        "arkadaşlar", "şimdi", "burada", "gördüğünüz gibi", 
        "yani", "efendim", "öncelikle", "şeklinde", "tabii ki",
        "mesela", "örneğin", "bu noktada", "aslında", "muhakkak",
        "baktığımızda", "diyebiliriz", "açıkçası", "zaten"
    ]
    
    for filler in fillers:
        pattern = re.compile(r'\b' + re.escape(filler) + r'\b', re.IGNORECASE)
        text = pattern.sub('', text)
        
    # 3. VERB CONJUGATION TRANSFORMATION (The Core Logic)
    # Target: Convert "yapıyoruz" (we are doing) to "yapınız" (do it / imperative formal)
    # or "seçiyoruz" -> "seçiniz"
    
    # Common Present Continuous Suffixes in Turkish:
    # -iyor, -ıyor, -uyor, -üyor
    # 1st Plural: -iyoruz, -ıyoruz, -uyoruz, -üyoruz
    # 2nd Plural: -iyorsunuz, ...
    # Future: -acağız, -eceğiz
    
    # Rules:
    # (word root) + (optional neg) + (suffix)
    
    # Map: "iyoruz" -> "iniz", "ıyoruz" -> "ınız", "uyoruz" -> "unuz", "üyoruz" -> "ünüz"
    
    replacements = [
        (r'(\w+)iyoruz\b', r'\1iniz'),
        (r'(\w+)ıyoruz\b', r'\1ınız'),
        (r'(\w+)uyoruz\b', r'\1unuz'),
        (r'(\w+)üyoruz\b', r'\1ünüz'),
        
        # Second person plural nuances
        (r'(\w+)iyorsunuz\b', r'\1iniz'),
        (r'(\w+)ıyorsunuz\b', r'\1ınız'),
        (r'(\w+)uyorsunuz\b', r'\1unuz'),
        (r'(\w+)üyorsunuz\b', r'\1ünüz'),
        
        # Basic imperative softener "lütfen" removal if valid, but keeping it is polite? 
        # Actually corporate docs rarely use "lütfen".
        
        # Future tense -> Passive/Formal Future
        # e.g. "yapacağız" -> "yapılacaktır" (hard to automate perfectly without root analysis)
        # fallback: "yapacağız" -> "yapınız" might be safer for instructions
        (r'(\w+)eceğiz\b', r'\1iniz'),
        (r'(\w+)acağız\b', r'\1ınız'),
        
        # "ediyoruz" -> "ediniz" is handled by above generic rules but "etmek" requires care?
        # "ediyoruz" -> "ed"+"iniz" works.
        
        # "basıyoruz" -> "basınız" works.
        
        # "girelim" -> "giriniz"
        (r'(\w+)elim\b', r'\1iniz'),
        (r'(\w+)alım\b', r'\1ınız'),
    ]
    
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text, flags=re.IGNORECASE)

    # 4. Clean up formatting
    # Fix punctuation spacing
    text = re.sub(r'\s+([.,!?;:])', r'\1', text)
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Capitalize sentences
    if text:
        # Split by sentence terminators to capitalize each sentence
        sentences = re.split(r'([.!?]+)', text)
        cleaned_sentences = []
        for s in sentences:
            s_clean = s.strip()
            if s_clean and not re.match(r'^[.!?]+$', s_clean):
                s_clean = s_clean[0].upper() + s_clean[1:]
                cleaned_sentences.append(s_clean)
            elif re.match(r'^[.!?]+$', s_clean):
                if cleaned_sentences:
                    cleaned_sentences[-1] += s_clean
                    
        text = " ".join(cleaned_sentences)
        
    return text

def create_document(video_path, vtt_path, output_docx):
    # No API Key needed
    
    # 1. Parse Subtitles
    print(f"Parsing {vtt_path}...")
    captions = []
    for caption in webvtt.read(vtt_path):
        text = caption.text.strip().replace('\n', ' ')
        start_sec = time_string_to_seconds(caption.start)
        captions.append({
            'start': start_sec,
            'text': text
        })

    # 2. Scene Detection Logic
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        print("Error opening video")
        return

    img_dir = "extracted_images"
    if not os.path.exists(img_dir):
        os.makedirs(img_dir)

    print("Processing video and text (Smart Mode)...")
    
    final_sections = []
    
    last_saved_img_path = None
    current_section = {
        'text_buffer': [],
        'image_path': None
    }
    
    chunk_start_time = 0
    
    for i, cap_data in enumerate(captions):
        text = clean_text_with_regex(cap_data['text'])
        if not text:
            continue
            
        time = cap_data['start']
        
        # Check current frame
        ret, frame = get_frame(cap, time)
        if not ret:
            current_section['text_buffer'].append(text)
            continue
            
        temp_img_name = "temp_current_frame.jpg"
        cv2.imwrite(temp_img_name, frame)
        
        diff_score = 0
        if last_saved_img_path:
            diff_score = calculate_image_difference(last_saved_img_path, temp_img_name)
        else:
            diff_score = 100.0 
            
        # Threshold (Scene Change)
        SCENE_CHANGE_THRESHOLD = 5.0 
        time_diff = time - chunk_start_time
        
        is_scene_change = (diff_score > SCENE_CHANGE_THRESHOLD) and (time_diff > 4.0)
        
        if is_scene_change or i == 0:
            # Save PREVIOUS
            if current_section['text_buffer'] or current_section['image_path']:
                current_section['final_text'] = " ".join(current_section['text_buffer'])
                final_sections.append(current_section)
            
            # Start NEW
            new_img_filename = f"frame_{int(time)}.jpg"
            new_img_path = os.path.join(img_dir, new_img_filename)
            import shutil
            shutil.move(temp_img_name, new_img_path)
            
            last_saved_img_path = new_img_path
            chunk_start_time = time
            
            current_section = {
                'text_buffer': [text],
                'image_path': new_img_path
            }
            print(f"Scene Change detected at {time}s (Diff: {diff_score:.2f}).")
            
        else:
            # Append text
            current_section['text_buffer'].append(text)

    # Append last
    if current_section['text_buffer']:
         current_section['final_text'] = " ".join(current_section['text_buffer'])
         final_sections.append(current_section)

    cap.release()
    if os.path.exists("temp_current_frame.jpg"):
        os.remove("temp_current_frame.jpg")

    # 3. Generate Word Doc (Same as before)

    print("Generating Word Document...")
    doc = Document()
    
    # Styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Cover
    doc.add_heading('Kullanıcı Dokümanı', 0)
    p = doc.add_paragraph('Otomatik Oluşturulmuş Doküman v2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    
    # Process Sections
    doc.add_heading('Süreç Adımları', level=1)
    
    for idx, section in enumerate(final_sections):
        if not section.get('final_text'):
            continue
            
        step_num = idx + 1
        
        # Step Header
        h = doc.add_heading(f"Adım {step_num}", level=2)
        
        # Text
        doc.add_paragraph(section['final_text'])
        
        # Image
        if section['image_path']:
            try:
                doc.add_picture(section['image_path'], width=Inches(6))
                cap = doc.add_paragraph(f"Ekran Görüntüsü {step_num}")
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cap.style = "Caption"
            except Exception as e:
                print(f"Image add error: {e}")
                
        doc.add_paragraph("") # Space
    
    doc.save(output_docx)
    print(f"Saved {output_docx}")

if __name__ == "__main__":
    # Hardcoded paths based on user files
    video_file = r"..\1-El Terminali Eğitimi-20250623_092431-Toplantı Kaydı.mp4"
    vtt_file = r"..\1-El-Terminali-E-itimi-20250623-092431-Toplant-Kayd-mp4-6e84e604-09f0.vtt"
    output_file = r"..\Taslak_Dokuman_v4.docx"
    
    create_document(video_file, vtt_file, output_file)
