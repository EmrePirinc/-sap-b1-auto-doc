
from docx import Document
from docx.oxml.document import CT_Body

def clean_template(path, output_path):
    doc = Document(path)
    
    print(f"Original Paragraphs: {len(doc.paragraphs)}")
    
    # We want to keep everything up to the first Table (History Table)
    # and maybe a few paragraphs after it if they are empty spacers?
    # Or simpler: Find "1. Amaç" and delete everything from there?
    
    start_delete = False
    to_delete = []
    
    # Method 1: Iterate XML body children (p and tbl)
    body = doc.element.body
    
    # We need to find the "History Table" index
    # Assuming it's the first table
    if len(doc.tables) == 0:
        print("No tables found!")
        return
        
    history_table = doc.tables[0]._element
    
    found_table = False
    
    # Iterate over valid children of body (p and tbl)
    for child in body.iterchildren():
        if child == history_table:
            found_table = True
            continue
            
        if found_table:
            # Delete everything AFTER the history table? 
            # But wait, usually usually "1. Amaç" follows. 
            # Let's check what's between Table and "1. Amaç"
            pass
            
    # Revised Logic:
    # 1. Find the History Table.
    # 2. Iterate all siblings AFTER the History Table.
    # 3. Remove them.
    
    # Get index of history table
    try:
        table_index = body.index(history_table)
    except ValueError:
        print("Table not in body?")
        return

    print(f"History Table is at Body Index: {table_index}")
    
    # Slice children after table
    # Note: We cannot slice a dynamic list while iterating if we modify it, so collect first
    children_to_remove = []
    for i in range(table_index + 1, len(body)):
        children_to_remove.append(body[i])
        
    print(f"Removing {len(children_to_remove)} elements after the table.")
    
    for child in children_to_remove:
        body.remove(child)
        
    doc.save(output_path)
    print(f"Cleaned doc saved to {output_path}")

if __name__ == "__main__":
    clean_template(r"..\Çoklu Para Birimi Sihirbazı Kullanıcı Dokümanı.docx", r"..\Template_Cleaned_Test.docx")
