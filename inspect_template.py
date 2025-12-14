
from docx import Document

def inspect_structure(path):
    doc = Document(path)
    print(f"--- STRUCTURE OF {path} ---")
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    print("\n--- FIRST 20 PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs[:20]):
        print(f"Para {i} [Style: {p.style.name}]: {p.text[:50]}")
        
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}: {len(table.rows)} rows")
        # Print first cell content to identify
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            print(f"  First cell: {table.rows[0].cells[0].text[:50]}")

if __name__ == "__main__":
    inspect_structure(r"..\Çoklu Para Birimi Sihirbazı Kullanıcı Dokümanı.docx")
